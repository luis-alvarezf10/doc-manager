import flet as ft
from src.utils.colors import *
from src.app.ui.widgets.custom_app_bar import custom_app_bar
from docx import Document
import pandas as pd
import os
from datetime import datetime

def document_generation_view(page: ft.Page, back_callback=None):
    def handle_close(e):
        if back_callback:
            back_callback()
    
    app_bar = custom_app_bar(
        text="Gestión de Empleados",
        on_click=handle_close
    )
    
    # Lista para almacenar empleados
    empleados_data = []
    
    # Tabla de empleados
    empleados_table = ft.DataTable(
        columns=[
            ft.DataColumn(ft.Text("Nombre")),
            ft.DataColumn(ft.Text("Cédula")),
            ft.DataColumn(ft.Text("Cargo")),
            ft.DataColumn(ft.Text("Salario")),
            ft.DataColumn(ft.Text("Acciones"))
        ],
        rows=[]
    )
    
    status_text = ft.Text("", size=14)
    
    # Campos del formulario
    nombre_field = ft.TextField(label="Nombre Completo")
    cedula_field = ft.TextField(label="Cédula de Identidad")
    direccion_field = ft.TextField(label="Dirección")
    cargo_field = ft.TextField(label="Cargo")
    salario_field = ft.TextField(label="Salario (Bs.)")
    fecha_inicio_field = ft.TextField(label="Fecha de Inicio (DD/MM/AAAA)")
    
    # Datos de la empresa
    empresa_nombre_field = ft.TextField(label="Nombre de la Empresa")
    empresa_rif_field = ft.TextField(label="RIF de la Empresa")
    empresa_direccion_field = ft.TextField(label="Dirección de la Empresa")
    
    def agregar_empleado(e):
        if not all([nombre_field.value, cedula_field.value, cargo_field.value, salario_field.value]):
            status_text.value = "Completa todos los campos obligatorios"
            status_text.color = "red"
            page.update()
            return
        
        empleado = {
            "nombre": nombre_field.value,
            "cedula": cedula_field.value,
            "direccion": direccion_field.value,
            "cargo": cargo_field.value,
            "salario": salario_field.value,
            "fecha_inicio": fecha_inicio_field.value,
            "empresa_nombre": empresa_nombre_field.value,
            "empresa_rif": empresa_rif_field.value,
            "empresa_direccion": empresa_direccion_field.value
        }
        
        empleados_data.append(empleado)
        actualizar_tabla()
        limpiar_campos()
        
        status_text.value = f"Empleado agregado correctamente"
        status_text.color = "green"
        page.update()
    
    def eliminar_empleado(index):
        empleados_data.pop(index)
        actualizar_tabla()
        status_text.value = "Empleado eliminado"
        status_text.color = "blue"
        page.update()
    
    def actualizar_tabla():
        empleados_table.rows.clear()
        for i, emp in enumerate(empleados_data):
            empleados_table.rows.append(
                ft.DataRow(
                    cells=[
                        ft.DataCell(ft.Text(emp["nombre"])),
                        ft.DataCell(ft.Text(emp["cedula"])),
                        ft.DataCell(ft.Text(emp["cargo"])),
                        ft.DataCell(ft.Text(f"Bs. {emp['salario']}")),
                        ft.DataCell(
                            ft.IconButton(
                                icon=ft.Icons.DELETE,
                                icon_color="red",
                                on_click=lambda e, idx=i: eliminar_empleado(idx)
                            )
                        )
                    ]
                )
            )
        
        # Habilitar botones si hay empleados
        guardar_excel_btn.disabled = len(empleados_data) == 0
        generar_contratos_btn.disabled = len(empleados_data) == 0
        page.update()
    
    def limpiar_campos():
        nombre_field.value = ""
        cedula_field.value = ""
        direccion_field.value = ""
        cargo_field.value = ""
        salario_field.value = ""
        fecha_inicio_field.value = ""
    
    def guardar_excel(e):
        if not empleados_data:
            status_text.value = "No hay empleados para guardar"
            status_text.color = "red"
            page.update()
            return
        
        try:
            # Crear DataFrame
            df = pd.DataFrame(empleados_data)
            
            # Guardar Excel en PLAF_system
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
            plaf_system_path = os.path.join(desktop_path, "PLAF_system")
            output_dir = os.path.join(plaf_system_path, "Empleados_Excel")
            os.makedirs(output_dir, exist_ok=True)
            
            filename = f"Empleados_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            output_path = os.path.join(output_dir, filename)
            
            df.to_excel(output_path, index=False)
            
            status_text.value = f"Excel guardado: {output_path}"
            status_text.color = "green"
            
        except Exception as ex:
            status_text.value = f"Error al guardar Excel: {str(ex)}"
            status_text.color = "red"
        
        page.update()
    
    def generar_contratos(e):
        if not empleados_data:
            status_text.value = "No hay empleados para generar contratos"
            status_text.color = "red"
            page.update()
            return
        
        try:
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
            plaf_system_path = os.path.join(desktop_path, "PLAF_system")
            output_dir = os.path.join(plaf_system_path, "Contratos_Generados")
            os.makedirs(output_dir, exist_ok=True)
            
            contratos_generados = 0
            
            for emp in empleados_data:
                doc = Document()
                
                # Crear contrato LOTTT
                doc.add_heading('CONTRATO DE TRABAJO', 0)
                doc.add_paragraph(f"En la ciudad de Caracas, a los {datetime.now().strftime('%d días del mes de %B del año %Y')}, comparecen:")
                
                doc.add_paragraph(f"Por una parte: {emp.get('empresa_nombre', '[EMPRESA]')}, venezolana, mayor de edad, titular del RIF número {emp.get('empresa_rif', '[RIF]')}, domiciliada en {emp.get('empresa_direccion', '[DIRECCIÓN]')}, en su carácter de EMPLEADOR.")
                
                doc.add_paragraph(f"Por la otra parte: {emp['nombre']}, venezolano, mayor de edad, titular de la cédula de identidad número {emp['cedula']}, domiciliado en {emp.get('direccion', '[DIRECCIÓN]')}, en su carácter de TRABAJADOR.")
                
                doc.add_heading('CLÁUSULAS:', level=1)
                doc.add_paragraph(f"PRIMERA: El trabajador prestará sus servicios en el cargo de {emp['cargo']}, con un salario mensual de {emp['salario']} bolívares.")
                
                doc.add_paragraph(f"SEGUNDA: La relación laboral iniciará el {emp.get('fecha_inicio', '[FECHA]')}.")
                
                doc.add_paragraph("TERCERA: Este contrato se rige por las disposiciones de la Ley Orgánica del Trabajo, los Trabajadores y las Trabajadoras.")
                
                # Guardar contrato individual
                filename = f"Contrato_{emp['nombre'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
                output_path = os.path.join(output_dir, filename)
                doc.save(output_path)
                
                contratos_generados += 1
            
            status_text.value = f"{contratos_generados} contratos generados en: {output_dir}"
            status_text.color = "green"
            
        except Exception as ex:
            status_text.value = f"Error al generar contratos: {str(ex)}"
            status_text.color = "red"
        
        page.update()
    
    # Botones principales
    agregar_btn = ft.ElevatedButton(
        text="Agregar Empleado",
        bgcolor="green",
        color="white",
        icon=ft.Icons.PERSON_ADD,
        on_click=agregar_empleado
    )
    
    guardar_excel_btn = ft.ElevatedButton(
        text="Guardar Excel",
        bgcolor="blue",
        color="white",
        icon=ft.Icons.TABLE_VIEW,
        disabled=True,
        on_click=guardar_excel
    )
    
    generar_contratos_btn = ft.ElevatedButton(
        text="Generar Contratos",
        bgcolor="orange",
        color="white",
        icon=ft.Icons.DESCRIPTION,
        disabled=True,
        on_click=generar_contratos
    )
    
    return ft.Column(
        controls=[
            app_bar,
            ft.Column([
                ft.Text("Gestión de Empleados y Contratos", size=24, weight=ft.FontWeight.BOLD),
                
                # Datos de la empresa
                ft.Text("Datos de la Empresa", size=18, weight=ft.FontWeight.BOLD),
                ft.Row([empresa_nombre_field, empresa_rif_field], spacing=10),
                empresa_direccion_field,
                
                ft.Divider(),
                
                # Formulario de empleado
                ft.Text("Agregar Empleado", size=18, weight=ft.FontWeight.BOLD),
                ft.Row([nombre_field, cedula_field], spacing=10),
                ft.Row([cargo_field, salario_field], spacing=10),
                ft.Row([direccion_field, fecha_inicio_field], spacing=10),
                
                agregar_btn,
                
                ft.Divider(),
                
                # Tabla de empleados
                ft.Text("Empleados Registrados", size=18, weight=ft.FontWeight.BOLD),
                ft.Column([
                    empleados_table
                ], height=200, scroll=ft.ScrollMode.AUTO),
                
                # Botones de acción
                ft.Row([
                    guardar_excel_btn,
                    generar_contratos_btn
                ], spacing=10),
                
                status_text
            ], spacing=15, scroll=ft.ScrollMode.AUTO, expand=True)
        ],
        expand=True
    )