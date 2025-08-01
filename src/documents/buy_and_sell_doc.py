import flet as ft
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from datetime import datetime
from src.documents.formated_text import add_formatted_text

from src.utils.convert_to_letters import number_to_letters

def generate_buy_and_sell_doc(vendedor_data, comprador_data, inmueble_data, oficina_data, page, input_filename):
    """
    Genera un contrato de compraventa profesional con formatos especiales
    
    Args:
        vendedor_data (dict): Datos del vendedor
        comprador_data (dict): Datos del comprador  
        inmueble_data (dict): Datos del inmueble
        oficina_data (dict): Datos de la oficina
        page: Objeto página de Flet
    
    Returns:
        str: Ruta del archivo generado
    """
    if 'Precio_letras' not in inmueble_data or 'area' not in inmueble_data or 'fecha_legal' not in inmueble_data:
        # Conversión de precio
        if 'Precio' in inmueble_data:
            precio_str = str(inmueble_data['Precio'])
            precio_numerico = precio_str.replace('.', '').replace(',', '.')
            try:
                inmueble_data['Precio_letras'] = number_to_letters(float(precio_numerico))
            except ValueError:
                inmueble_data['Precio_letras'] = "[PRECIO INVÁLIDO]"
        
        # Conversión de área (solo parte entera para metros cuadrados)
        if 'Superficie (m²)' in inmueble_data:
            area_str = str(inmueble_data['Superficie (m²)'])
            area_numerica = area_str.replace(',', '.')
            try:
                # Tomamos solo la parte entera para metros cuadrados
                area_entera = int(float(area_numerica))
                inmueble_data['area'] = number_to_letters(area_entera, es_metros_cuadrados=True)
            except ValueError:
                inmueble_data['area'] = "[ÁREA INVÁLIDA]"
        
        # Fecha legal si no existe
        if 'fecha_legal' not in inmueble_data:
            inmueble_data['fecha_legal'] = datetime.now().strftime("%d/%m/%Y")
        
    try:
        # Función auxiliar para formato
      
        # Crear documento
        doc = Document()
        
        
        
        # Configuración de estilos
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(14)
        
        
        section = doc.sections[0]  # Acceder a la primera sección (normalmente hay solo una)
        section.page_width = Inches(8.5)    
        section.page_height = Inches(14)
        
        # Establecer márgenes (en pulgadas)
        section.left_margin = Cm(3)    
        section.right_margin = Cm(3)  
        section.top_margin = Cm(3)    
        section.bottom_margin = Cm(1.5)
        
        impre = doc.add_paragraph()
        runs = [
            ("PEDRO LUIS ALVAREZ", True),
            ("ABOGADO", True),
            ("I.P.S.A: 41.432", False)
        ]
        for text, add_break in runs:
            run = impre.add_run(text)
            run.font.size = Pt(9)
            run.font.name = "Algerian"
            if add_break:
                run.add_break()
        impre.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        intro_p = doc.add_paragraph()
        add_formatted_text(intro_p, f"Yo, {vendedor_data['Nombre'].upper()}, ", bold=True)
        add_formatted_text(intro_p, f"{vendedor_data['Nacionalidad']}, mayor de edad, {vendedor_data['Ocupación'].lower()}, ")
        add_formatted_text(intro_p, f"{vendedor_data['Estado civil'].lower()}, ")
        add_formatted_text(intro_p, "titular de la cédula de identidad personal No. ")
        if vendedor_data.get("Nacionalidad", "").lower() in  ["venezolano", "venezolana", "venezolano/a"]:
            add_formatted_text(intro_p, "V-", bold=True)
        else:
            add_formatted_text(intro_p, "E-", bold=True)
        add_formatted_text(intro_p, f"{vendedor_data.get('Cédula', 'No especificado')}, ", bold=True)
        add_formatted_text(intro_p, "con Registro de Información Fiscal (R.I.F) No. ")
        add_formatted_text(intro_p, f"{vendedor_data['RIF']}, ", bold=True)
        add_formatted_text(intro_p, f"con domicilio en esta ciudad de {vendedor_data['Domicilio (Ciudad)']}, ")
        add_formatted_text(intro_p, f"Municipio {vendedor_data['Domicilio (Municipio)']} ")
        add_formatted_text(intro_p, f"del Estado {vendedor_data['Domicilio (Estado)']}, ")
        add_formatted_text(intro_p, "por medio de la presente declaro: Que doy en venta pura y simple, perfecta e irrevocable al ciudadano: ")
        add_formatted_text(intro_p, f"{comprador_data['Nombre'].upper()}, ", bold=True)
        add_formatted_text(intro_p, f"{comprador_data['Nacionalidad']}, mayor de edad, titular de la cédula de identidad No. ")
        add_formatted_text(intro_p, f"V-{comprador_data['Cédula']}, ", bold=True)
        add_formatted_text(intro_p, f"Estado Civil {comprador_data['Estado civil']}, con Registro de Información Fiscal (R.I.F) No. ")
        add_formatted_text(intro_p, f"{comprador_data['RIF']}, ", bold=True)
        add_formatted_text(intro_p, f"con domicilio en la ciudad de {comprador_data['Domicilio (Ciudad)']}, ")
        add_formatted_text(intro_p, f"Municipio {comprador_data['Domicilio (Municipio)']} ", bold=True)
        add_formatted_text(intro_p, f"del estado {comprador_data['Domicilio (Estado)']}, ")
        add_formatted_text(intro_p, "un (01) bien inmueble de mi legítima propiedad, constituido por ")
        add_formatted_text(intro_p, f"un(a) (01) {inmueble_data['Tipo de Inmueble'].lower()} ubicado en {inmueble_data['Ubicación'].lower()}", underline=True)
        add_formatted_text(intro_p, f", de la ciudad de {inmueble_data['Domicilio (Ciudad)']}, Parroquia {inmueble_data['Domicilio (Parroquia)']} Jurisdicción del Municipio {inmueble_data['Domicilio (Municipio)']} del Estado {inmueble_data['Domicilio (Estado)']}", underline=True)
        add_formatted_text(intro_p, f", distinguido con el ")
        add_formatted_text(intro_p, f"CÓDIGO CATASTRAL: {inmueble_data['Código catastral']},", bold=True, underline=True)
        add_formatted_text(intro_p, f" en clavada en un área de terreno que no forma parte de esta venta, constante de una superficie de ")
        add_formatted_text(intro_p, f"{inmueble_data['area'].upper()} ({inmueble_data['Superficie (m²)']} m²),", bold=True)
        add_formatted_text(intro_p, " con los siguientes linderos: ")
        add_formatted_text(intro_p, "NORTE", bold=True, underline=True)
        add_formatted_text(intro_p, f": {inmueble_data['Límite Norte']}; ")
        add_formatted_text(intro_p, "SUR:", bold=True, underline=True)
        add_formatted_text(intro_p, f" {inmueble_data['Límite Sur']}; ")
        add_formatted_text(intro_p, "ESTE:", bold=True, underline=True)
        add_formatted_text(intro_p, f" {inmueble_data['Límite Este']}; ")
        add_formatted_text(intro_p, "OESTE:", bold=True, underline=True)
        add_formatted_text(intro_p, f" {inmueble_data['Límite Oeste']}; ")
        add_formatted_text(intro_p, "El precio de venta establecido entre las partes para la mencionada negociación es la cantidad de ")
        add_formatted_text(intro_p, f"{inmueble_data['Precio_letras'].upper()} (Bs. {inmueble_data['Precio']}),", bold=True)
        add_formatted_text(intro_p, " cantidad de dinero que declaro recibir en este acto mediante cheque No. ")
        add_formatted_text(intro_p, f"{inmueble_data['Número de cheque']}, ", bold=True)
        add_formatted_text(intro_p, "librado contra la cuenta corriente No. ")
        add_formatted_text(intro_p, f"{inmueble_data['Cuenta a depositar']}, ", bold=True)
        add_formatted_text(intro_p, f"correspondiente al BANCO {inmueble_data['Banco'].upper()}, a mi entera y cabal satisfacción.")
        add_formatted_text(intro_p, f" El inmueble objeto de la presente venta me pertenece  conforme a Documento debidamente protocolizado por   ante   la   Oficina   de Registro  Público  del  Municipio {oficina_data['Domicilio (Municipio)']}  del  Estado  {oficina_data['Domicilio (Estado)']}, de fecha")
        add_formatted_text(intro_p, f" {oficina_data['Fecha_legal']} INSCRITO BAJO EL NUMERO DE {oficina_data['Número de folio']} PROTOCOLO {oficina_data['Protocolo']}, TOMO {oficina_data['Tomo']}, {oficina_data['Trimestre referido'].upper()} TRIMESTRE DEL REFERIDO AÑO {oficina_data['Ano_documento'].upper()}.", bold=True, underline=True)
        add_formatted_text(intro_p, f" Con el otorgamiento de este documento hago la tradición legal del inmueble aquí vendido y lo coloco en posesión y propiedad del mismo, obligándome al saneamiento de Ley; asimismo declaro que sobre el mismo no existe gravamen alguno que lo afecte y nada adeuda por concepto de impuestos Nacionales, Estadales o Municipales. Y yo,"),
        add_formatted_text(intro_p, f" {vendedor_data['Nombre'].upper()}, ", bold=True)
        add_formatted_text(intro_p, f"plenamente identificado(a), acepto la venta que se me hace en los términos antes expuestos. En esta ciudad de {oficina_data['Domicilio (Ciudad)']}, Municipio {oficina_data['Domicilio (Municipio)']} del Estado {oficina_data['Domicilio (Estado)']}, a la fecha de su otorgamiento.")
        
        intro_p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        intro_p.paragraph_format.first_line_indent = Pt(28)
        # Guardar documento
        documents_path = os.path.join(os.path.expanduser("~"), "Documents")
        manager_path = os.path.join(documents_path, "Axiology Document Manager")
        contracts_path = os.path.join(manager_path, "Contratos Generados")
        sales_path = os.path.join(contracts_path, "Contratos de Compra Venta")
        os.makedirs(sales_path, exist_ok=True)
        
        
        if input_filename == "":
            filename =  f"CONTRATO DE COMPRAVENTA DE PROPIEDAD {vendedor_data['Nombre'].upper()} Y {comprador_data['Nombre'].upper()}.docx"
        else:
            filename = f"{input_filename.upper()}.docx"
            
        full_path = os.path.join(sales_path, filename)
        doc.save(full_path)

        # Notificación de éxito
        if page:
            page.update()

        return full_path

    except Exception as ex:
        # notificacion de fallo
        if page:
            print(f"Error: {str.ex}")
            page.update()
