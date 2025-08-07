import flet as ft
from docx import Document
from docx.shared import Pt, Inches, Cm
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from src.documents.formated_text import add_formatted_text
from src.documents.add_firm import add_firm
from src.utils.normalizar import es_venezolano, lista_con_y, normalizar_nacionalidad, normalizar_ocupaciones, normalizar_estados_civiles, imprimir_domicilios_unicos, normalizar_acciones, normalizar_cargos, pluralizar_cargo
from src.utils.convert_to_letters import number_to_letters


def generate_constitutive_act(accionistas: list, presentante: dict, empresa: dict, total_acciones, comisario: dict, input_filename:str):
    
    if 'Capital_letras' not in empresa:
        # Conversión de capital
        if 'Capital' in empresa:
            capital_str = str(empresa['Capital'])
            capital_numerico = capital_str.replace('.', '').replace(',', '.')
            try:
                empresa['Capital_letras'] = number_to_letters(float(capital_numerico))
            except ValueError:
                empresa['Capital_letras'] = "[Error, convert num to str]"
                
    n_acciones = str(total_acciones)
    n_acciones_numerico = n_acciones.replace('.', '').replace(', ', '.')
    n_acciones_letras = ""
    try:
        n_acciones_letras = number_to_letters(float(n_acciones_numerico), incluir_moneda=False)
    except ValueError:
        n_acciones_letras = "[Error, convert num to str]"
        
    # Calcular valor nominal por acción
    try:
        capital_numerico = float(str(empresa['Capital']).replace('.', '').replace(',', '.'))
        total_acciones_num = int(total_acciones)
        nomina = capital_numerico / total_acciones_num
        n_nomina_letras = number_to_letters(nomina)
    except (ValueError, ZeroDivisionError):
        nomina = 0
        n_nomina_letras = "[Error al calcular]"
    try:
        doc = Document()
        
        style = doc.styles['Normal']
        style.font.name = 'Calibri' 
        style.font.size = Pt(14)
        
        section = doc.sections[0] 
        section.page_width = Inches(8.5)
        section.page_height = Inches(14)
        
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)  
        section.top_margin = Cm(3)    
        section.bottom_margin = Cm(1.5)

        add_firm(doc)
        
        title = doc.add_paragraph()
        add_formatted_text(title, "CIUDADANO:\nREGISTRADOR MERCANTIL PRIMERO DE LA CIRCUNSCRIPCIÓN JUDICIAL DEL ESTADO ANZOATEGUI.\nSU DESPACHO:", bold=True)
        
        
        intro_p = doc.add_paragraph()
        add_formatted_text(intro_p, f"Yo, {presentante["nombre"].upper()}, ", bold=True)
        add_formatted_text(intro_p, f"{presentante["nacionalidad"]}, mayor de edad, {presentante["estado civil"]}, {presentante["ocupacion"]}, titular de la cédula de identidad personal No. {presentante["cedula"]}, con Registro de Información Fiscal (R.I.F) No. {presentante["rif"]} con domicilio en esta ciudad de {presentante["ciudad"]}, Municipio {presentante["municipio"]}, del Estado {presentante["estado"]}, debidamente")
        if presentante["nacionalidad"] == "venezolano":
            autorizado_text = "autorizado"
        else:
            autorizado_text = "autorizada"
        add_formatted_text(intro_p, f" {autorizado_text} en la presente acta de la compañía ")
        add_formatted_text(intro_p, f"\"{empresa["Nombre"].upper()}\"", bold=True)
        add_formatted_text(intro_p, " empresa de este domicilio; Ante usted ocurro de conformidad con lo establecido en el Artículo 215 del Código de Comercio Vigente, para presentar el Documento Constitutivo de mi representada, el cual ha sido redactado con suficiente amplitud para que sirva de Estatutos de la misma. Participación que hago a Usted, a los fines de que, una vez cumplidos los trámites de inscripción, se sirva ordenar su Registro, Fijación y Publicación, con el ruego de expedirme sendas copias certificadas, de la presente participación del Documento Constitutivo, y del auto correspondiente. Es Justicia, que espero en esta ciudad a la fecha de su presentación. ")
        intro_p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        
        firm = doc.add_paragraph()
        add_formatted_text(firm, "\n\n\n", bold=True)
        add_formatted_text(firm, f"\t\t\t\t\t\t{presentante["nombre"].upper()}\n\t\t\t\t\t\t\tC.I. {presentante["cedula"]}", bold=True)
        run = firm.add_run()
        run.add_break(WD_BREAK.PAGE)
        
        nombres_lista = [a["Nombre"] for a in accionistas]

        cedulas_lista = [
            f"{'V' if es_venezolano(a['Nacionalidad']) else 'E'}-{a['Cédula']}"
            for a in accionistas
        ]

        rif_lista = [
            a["RIF"] if a["RIF"].startswith("J-") else f"J-{a['RIF']}"
            for a in accionistas
        ]# evita duplicar J-

        
        nombres = lista_con_y(nombres_lista)
        cedulas = lista_con_y(cedulas_lista)
        rif = lista_con_y(rif_lista)
        nacionalidades = normalizar_nacionalidad(accionistas)
        ocupaciones = normalizar_ocupaciones(accionistas)   
        estado = normalizar_estados_civiles(accionistas)
        domicilio = imprimir_domicilios_unicos(accionistas)
        # impre
        add_firm(doc)
        title2 = doc.add_paragraph()
        add_formatted_text(title2, f"DOCUMENTO CONSTITUTIVO EMPRESA\n\"{empresa["Nombre"].upper()}\"", bold=True)
        title2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        
        text_2 = doc.add_paragraph()
        add_formatted_text(text_2, " Nosotros, ")
        add_formatted_text(text_2, f"{nombres.upper()}", bold=True)
        add_formatted_text(text_2, f", {nacionalidades}, mayores de edad, {ocupaciones}, de estado civil {estado}, titulares de la cédula de identidad personal numeros: {cedulas}, con Registro de Información Fiscal (R.I.F) número: {rif}, respectivamente con domicilio en {domicilio}; por medio del presente documento declaramos: Que hemos convenido en constituir como en efecto constituimos una Compañía Anónima, la cual se regirá por las Cláusulas contenidas en el presente documento:")
        text_2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        title_cap1 = doc.add_paragraph()
        add_formatted_text(title_cap1, "CAPÍTULO I:\nDENOMINACIÓN, OBJETO, DOMICICLIO Y DURACIÓN.", bold=True)
        title_cap1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        cap1 = doc.add_paragraph()
        add_formatted_text(cap1, "PRIMERA: ", bold=True)
        add_formatted_text(cap1, f"la compañía se denominará \"{empresa["Nombre"].upper()}\".\n")
        add_formatted_text(cap1, "SEGUNDA: ", bold=True)
        add_formatted_text(cap1, f"La compañía tendrá como objeto principal {empresa["Dedicación"]} y en fin cualquier actividad que esté relacionada o conexa con el objeto principal, dentro y fuera del país. ")
        add_formatted_text(cap1, "TERCERA: ", bold=True)
        add_formatted_text(cap1, "La compañía tendrá domicilio en ")
        add_formatted_text(cap1, f"{empresa['Ubicación']} en la ciudad de {empresa["Domicilio (Ciudad)"]}, Municipio {empresa["Domicilio (Municipio)"]} del Estado {empresa["Domicilio (Estado)"]}.".upper())
        add_formatted_text(cap1, " Pudiendo establecer Agencias, Oficinas y/o sucursales en cualquier lugar de la República Bolivariana de Venezuela o el exterior. ")
        add_formatted_text(cap1, "CUARTA: ", bold=True)
        add_formatted_text(cap1, f"La duración de la compañía será de ({empresa["Duración"]}) años, contados a partir de la fecha de su Inscripción por ante la Oficina de Registro Mercantil competente.")
        cap1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        title_cap2 = doc.add_paragraph()
        add_formatted_text(title_cap2, "CAPÍTULO II:\nDEL CAPITAL Y LAS ACCIONESSS", bold=True)
        title_cap2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        cap2 = doc.add_paragraph()
        add_formatted_text(cap2, "QUINTA: ", bold=True)
        add_formatted_text(cap2, "El capital de la Compañía es de")    
        add_formatted_text(cap2, f" {empresa["Capital_letras"].upper()} (Bs {empresa["Capital"]}) ", bold=True)
        add_formatted_text(cap2, f"dividido entre {n_acciones_letras} ({total_acciones}) acciones con un valor nominal de {n_nomina_letras} (Bs. {nomina:.2f}) cada una el cual ha sido totalmente suscrito y pagado, así: ")
        
        # Agregar descripción de acciones por accionista
        acciones_texto = normalizar_acciones(accionistas, nomina)
        add_formatted_text(cap2, f"{acciones_texto}. Este capital ha sido pagado en su totalidad mediante inventario de bienes que se anexa. ")
        add_formatted_text(cap2, "SEXTA: ", bold=True)
        add_formatted_text(cap2, "Las acciones son iguales entre si y confieren a sus tenedores iguales derechos y obligaciones, son indivisibles respecto a la Compañía, la cual reconocerá a un solo propietario por cada acción y dan derecho a un voto en la Asamblea General de Accionistas.")
        cap2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        title_cap3 = doc.add_paragraph()
        add_formatted_text(title_cap3, "CAPÍTULO III:\nDE LA ADDMINISTRACIÓN", bold=True)
        title_cap3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # SÉPTIMA: Administración
        cargos_info = normalizar_cargos(accionistas)
        cap3 = doc.add_paragraph()
        add_formatted_text(cap3, "SÉPTIMA: ", bold=True)
        add_formatted_text(cap3, "La Administración de la Compañía estará a cargo de una Junta Directiva integrada por ")
        add_formatted_text(cap3, f"{cargos_info['cantidad_letras']} ({cargos_info['cantidad']}) ", bold=True)
        add_formatted_text(cap3, f"miembros que se denominarán: {cargos_info['cargo_principal'].upper()}, los cuales serán designados por la Asamblea General de Accionistas, durarán ")
        try:
            anos_letras = number_to_letters(int(empresa['Años de función de accionistas']), incluir_moneda=False)
        except:
            anos_letras = empresa['Años de función de accionistas']
        add_formatted_text(cap3, f"{anos_letras} ({empresa['Años de función de accionistas']}) ", bold=True)
        add_formatted_text(cap3, "años en sus funciones y  podrán  ser  reelectos,  depositaran  una  (1)  acción  como garantía de su gestión, de acuerdo a lo dispuesto en el Artículo 244 del Código de Comercio Venezolano vigente. ")
        add_formatted_text(cap3, "OCTAVA: ", bold=True)
        add_formatted_text(cap3, f"LOS {cargos_info['cargo_principal'].upper()} ACTUANDO CONJUNTA Y/O SEPARADAMENTE, tendrán las mayores facultades de administración y disposición, entre ellas las siguientes: 1.-Representar judicial y extrajudicialmente a la Compañía en todos los asuntos que se le presenten con facultades para demandar, contestar demandas y reconvenciones, convenir, desistir, transigir, comprometer en árbitros,   2.-   vender   bienes   muebles   e   inmuebles   propiedad   de   la  Empresa.3.- comprar bienes, enajenar y gravar, 4.- constituir Apoderados judiciales en Abogados de su confianza, fijándoles  sus  atribuciones  y  remuneraciones,  5.-abrir  y  cerrar  cuentas  bancarias, endosar, librar, avalar, cancelar, aceptar o emitir letras de cambio, cheques, pagares, 6.-otorgar poderes de cualquier tipo, o cualquier otro tipo de efecto de comercio, 7.-solicitar créditos. 8.- convocar y presidir la Asamblea de Accionistas, 9.-celebrar a nombre de la Compañía toda clase de contratos.10.-contratar personal y fijar su remuneración. ")
        cap3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        title_cap4 = doc.add_paragraph()
        add_formatted_text(title_cap4, "CAPÍTULO IV:\nDE LAS ASAMBLEAS", bold=True)
        title_cap4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        cap4 = doc.add_paragraph()
        add_formatted_text(cap4, "NOVENA: ", bold=True)
        add_formatted_text(cap4, ": Las Asambleas Ordinarias y Extraordinarias de la Compañía son la Suprema autoridad de la misma y tienen las facultades que le atribuyen las Leyes Nacionales que regulan la materia y el presente documento constitutivo. La Asamblea ordinaria se reunirá dentro de los noventa (90) días siguientes a la fecha de cierre del respectivo Ejercicio Económico. Las Asambleas Extraordinarias se reunirán cuando lo consideren conveniente los miembros de la Junta Directiva o a solicitud de un número de accionistas que representen el CINCUENTA Y UN POR CIENTO (51%) del Capital Social, para su validez y aprobación tanto ordinaria como extraordinaria. ")
        add_formatted_text(cap4, "DECIMA: ", bold=True)
        add_formatted_text(cap4, "Las convocatorias para las Asambleas sean Ordinarias o Extraordinarias deberán ser publicadas en un periodo de circulación regional con cinco (5) días de anticipación por lo menos al fijado para su reunión y de forma personal privado, según lo establecido en los  Artículo 277 y 279 del Código de Comercio. Los acuerdos de la Asamblea se tomarán por mayoría absoluta de los votos presentes, no pudiendo ésta constituirse sin que esté representado el cincuenta y un por ciento (51%) del Capital Social. Cuando se encuentren presentes todos los accionistas, podrá prescindirse de las formalidades de convocatoria establecidas por el Código de Comercio.")
        cap4.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        title_cap5 = doc.add_paragraph()
        add_formatted_text(title_cap5, "CAPÍTULO V:\nDEL EJERCICIO ECONÓMICO Y LAS CUENTAS", bold=True)
        title_cap5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        
        ## corregir aqui redaccion es necesario
        cap5 = doc.add_paragraph()
        add_formatted_text(cap5, "DÉCIMA PRIMERA: ", bold=True)
        add_formatted_text(cap5, ": El año Económico de la Compañía comenzará el 1º de Enero y terminará el 31 de Diciembre de cada año. En esta última fecha se practicará un Inventario y Balance General con el Estado Sumario de la situación activa y pasiva de la  Compañía  en  el  año  que termina de conformidad con lo establecido en el Código de Comercio. ")
        add_formatted_text(cap5, "DÉCIMA SEGUNDA: ", bold=True)
        add_formatted_text(cap5, "El primer ejercicio comenzará a partir de la Inscripción del presente documento en el Registro Mercantil y finalizará el treinta y uno (31) de Diciembre del mismo año. ")
        add_formatted_text(cap5, "DÉCIMA TERCERA: ", bold=True)
        add_formatted_text(cap5, "Al cierre de cada ejercicio anual una vez deducidos los gastos, hechas las amortizaciones y constituidos los fondos de previsión que se juzgare necesario establecer, de las utilidades se deducirá un Cinco por ciento (5%) como fondo de reserva hasta que dicho fondo alcance a un diez por ciento (10%) del Capital Social de la Compañía, el remanente o utilidades serán repartidas entre los Accionistas en proporción al valor nominal de sus respectivas acciones. ")
        add_formatted_text(cap5, "DÉCIMA CUARTA: ", bold=True)
        add_formatted_text(cap5, "Si llegare el caso de la liquidación de la Compañía, la Asamblea de Accionistas estará investida de los más amplios poderes para todo lo relacionado con ello y se regirá por lo pautado en el Código de Comercio. ")
        add_formatted_text(cap5, "DÉCIMA QUINTA: ", bold=True)
        add_formatted_text(cap5, "Todo lo no previsto en el presente Documento Constitutivo y Estatutos Sociales, se regirá por las disposiciones establecidas en el Código de Comercio vigente y en las Leyes Especiales que regulan la materia.")
        cap5.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        title_cap6 = doc.add_paragraph()
        add_formatted_text(title_cap6, "CAPÍTULO VI:\nDEL COMISARIO", bold=True)
        title_cap6.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        cap6 = doc.add_paragraph()
        add_formatted_text(cap6, "DÉCIMA SEXTA: ", bold=True)
        add_formatted_text(cap6, f"La Compañía tendrá un comisario que será elegido por la Asamblea General de Accionistas y durará cinco ({comisario['Años de función']}) años en el ejercicio de sus funciones, pudiendo ser reelegido.")
        
        title_cap7 = doc.add_paragraph()
        add_formatted_text(title_cap7, "CAPÍTULO VIII:\nDISPOSICIONES TRANSITORIAS", bold=True)
        title_cap7.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        cap7 = doc.add_paragraph()
        add_formatted_text(cap7, "DÉCIMA SÉPTIMA: ", bold=True)
        add_formatted_text(cap7, "Para ocupar los cargos de la primera Junta Directiva, la cual iniciara  sus funciones  a  partir  de    la inscripción  del   presente Documento en   el Registro Mercantil, han sido nombrados como ")
        
        # Normalizar cargos según la variedad
        cargos_unicos = set(a['Cargo'] for a in accionistas)
        
        # Si hay cargos directivos (Presidente, Vicepresidente, etc.), usar "MIEMBROS DE JUNTA DIRECTIVA"
        cargos_directivos = {"presidente", "vicepresidente", "secretario", "tesorero"}
        if any(cargo.lower() in cargos_directivos for cargo in cargos_unicos):
            cargo_plural = "MIEMBROS DE JUNTA DIRECTIVA"
        elif len(cargos_unicos) == 1:
            cargo_unico = list(cargos_unicos)[0]
            cargo_plural = pluralizar_cargo(cargo_unico)
        else:
            cargo_plural = "DIRECTORES GENERALES"  # Fallback genérico
        
        add_formatted_text(cap7, f"{cargo_plural.upper()}: ", bold=True)
        add_formatted_text(cap7, f"{nombres.upper()}. ", bold=True)
        add_formatted_text(cap7, "Fue designado como  Comisario a ")
        if comisario['Sexo'] == "Femenino":
            text_gen_licenced = "la licenciada"
        else:
            text_gen_licenced = "el licenciado"
        add_formatted_text(cap7, f"{text_gen_licenced} ")
        add_formatted_text(cap7, f"{comisario['Nombre'].upper()}, ", bold=True)
        add_formatted_text(cap7, f"{comisario['Nacionalidad']}, mayor de edad, titular de la cédula de identidad personal No. V-{comisario['Cédula']}, de este domicilio, e inscrita en el Colegio de Contadores bajo el No. {comisario['No de Colegio']}. ")
        if presentante['sexo'] == "Femenino":
            text_gen_pr = "autorizada la ciudadana"
        else:
            text_gen_pr = "autorizado el ciudadano"
        add_formatted_text(cap7, f"Queda suficientemente {text_gen_pr} ")
        add_formatted_text(cap7, f"{presentante['nombre'].upper()}, ", bold=True)
        add_formatted_text(cap7, f"{presentante['nacionalidad']}, mayor de edad, titular de la cédula de identidad personal No. V-{presentante['cedula']}, y de este domicilio, para firmar todo lo relacionado con la inscripción de este documento ante el Ciudadano Registrador Mercantil, de igual manera para tramitar y gestionar todo lo relativo a la obtención del Registro de Información Fiscal (R.I.F), por ante las autoridades del SENIAT, respectiva. En esta ciudad, a la fecha de su presentación.")
        cap7.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        # Firmas de los accionistas
        firmas = doc.add_paragraph()
        add_formatted_text(firmas, "\n\nFIRMAS DE LOS ACCIONISTAS:\n\n", bold=True)
        firmas.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Generar firmas para cada accionista
        for i, accionista in enumerate(accionistas):
            firma_accionista = doc.add_paragraph()
            add_formatted_text(firma_accionista, "\n_________________________________\n", bold=True)
            add_formatted_text(firma_accionista, f"{accionista['Nombre'].upper()}\n", bold=True)
            add_formatted_text(firma_accionista, f"C.I. {accionista['Cédula']}\n", bold=True)
            firma_accionista.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Agregar espacio entre firmas
            if i < len(accionistas) - 1:
                doc.add_paragraph("\n")
        

        
        
        documents_path = os.path.join(os.path.expanduser("~"), "Documents")
        manager_path = os.path.join(documents_path, "Axiology Document Manager")
        contracts_path = os.path.join(manager_path, "Contratos Generados")
        sales_path = os.path.join(contracts_path, "Actas Constitutivas")
        os.makedirs(sales_path, exist_ok=True)
        
        
        if input_filename == "":
            base = "ACTA CONSTITUTIVA"
            cont = 1
            filename =  f"{base} {empresa['Nombre'].upper()}.docx"
        
            while os.path.exists(os.path.join(sales_path, filename)):
                filename = f"{base} ({cont}).docx"
                cont += 1
        else:
            filename = f"{input_filename}.docx"
        
        
        full_path = os.path.join(sales_path, filename)
        doc.save(full_path)
        return full_path
    except Exception as e:
            print(f"Error al generar el documento: {e}")
            return None